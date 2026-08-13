from __future__ import annotations

import os
import re
import secrets
import shutil
import sys
import threading
import zipfile
from datetime import datetime, timedelta
from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.oxml.ns import qn
from flask import (
    Flask,
    abort,
    after_this_request,
    jsonify,
    render_template,
    request,
    send_file,
    url_for,
)

def get_asset_dir() -> Path:
    """Папка с шаблонами HTML, static и встроенным template.docx."""
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        return Path(sys._MEIPASS)  # type: ignore[attr-defined]
    return Path(__file__).resolve().parent


def get_data_dir() -> Path:
    """Папка для изменяемых файлов рядом с exe или рядом с исходниками."""
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return get_asset_dir()


ASSET_DIR = get_asset_dir()
DATA_DIR = get_data_dir()
BASE_DIR = DATA_DIR
UPLOAD_DIR = DATA_DIR / "uploads"
OUTPUT_DIR = DATA_DIR / "output"
TEMPLATE_PATH = UPLOAD_DIR / "template.docx"
BUNDLED_TEMPLATE_PATH = ASSET_DIR / "uploads" / "template.docx"
OFFER_TEMPLATE_PATH = UPLOAD_DIR / "template_offer.docx"
BUNDLED_OFFER_TEMPLATE_PATH = ASSET_DIR / "uploads" / "template_offer.docx"

FIELD_KEYS = [
    "FIO",
    "PASSPORT",
    "INN_EXEC",
    "ADDRESS_EXEC",
    "ORG_NAME",
    "INN_ORG",
    "ADDRESS_ORG",
    "OWNER_FIO",
    "PLACE",
    "SERVICES",
    "DATE",
    "VOLUME",
    "PRICE",
    "SIGN_DATETIME",
    "VALID_FROM",
    "VALID_TO",
    # Поля оферты (уникальные):
    "OFFER_SENT_DATE",
    "EXEC_PHONE",
    "EXEC_EMAIL",
    "EXEC_ACCBANK",
    "SIGN_DATETIMEOFF",
]

# Старые фиксированные номера сертификатов из шаблона. Они заменяются на новые
# случайные значения при каждой генерации, даже если пользователь загрузил старый
# шаблон, где эти номера всё ещё прописаны обычным текстом.
DEFAULT_ORG_CERT = "d733cea5-3307-48c4-8cca-4332615e08f8"
DEFAULT_EXEC_CERT = "2af19ff3-507f-4cc6-afd5-386881h37c9f"

app = Flask(
    __name__,
    template_folder=str(ASSET_DIR / "templates"),
    static_folder=str(ASSET_DIR / "static"),
)
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16 MB

UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

# Папка output/ — только для временных ZIP-архивов. При запуске очищаем её,
# чтобы она не копилась и не смущала пользователя (плюс убираем остатки
# незавершённых генераций после сбоя).
try:
    for _stale in OUTPUT_DIR.glob("*"):
        try:
            if _stale.is_dir():
                shutil.rmtree(_stale, ignore_errors=True)
            else:
                _stale.unlink(missing_ok=True)
        except OSError:
            pass
except OSError:
    pass

# В desktop/exe-версии исходный template.docx лежит внутри сборки,
# а рабочий шаблон должен храниться рядом с exe, чтобы /admin мог его менять.
if not TEMPLATE_PATH.exists() and BUNDLED_TEMPLATE_PATH.exists():
    try:
        if BUNDLED_TEMPLATE_PATH.resolve() != TEMPLATE_PATH.resolve():
            shutil.copy2(BUNDLED_TEMPLATE_PATH, TEMPLATE_PATH)
    except OSError:
        pass

if not OFFER_TEMPLATE_PATH.exists() and BUNDLED_OFFER_TEMPLATE_PATH.exists():
    try:
        if BUNDLED_OFFER_TEMPLATE_PATH.resolve() != OFFER_TEMPLATE_PATH.resolve():
            shutil.copy2(BUNDLED_OFFER_TEMPLATE_PATH, OFFER_TEMPLATE_PATH)
    except OSError:
        pass

# Microsoft Word через COM лучше запускать по одному процессу за раз.
# Это особенно важно при массовой генерации и параллельных запросах.
PDF_CONVERSION_LOCK = threading.Lock()


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/admin")
def admin():
    return render_template(
        "admin.html",
        template_exists=TEMPLATE_PATH.exists(),
        template_offer_exists=OFFER_TEMPLATE_PATH.exists(),
    )


@app.route("/admin/upload", methods=["POST"])
def upload_template():
    if "template" not in request.files:
        return jsonify(ok=False, error="Файл шаблона не передан."), 400

    file = request.files["template"]
    if not file or file.filename == "":
        return jsonify(ok=False, error="Файл не выбран."), 400

    if not file.filename.lower().endswith(".docx"):
        return jsonify(ok=False, error="Можно загрузить только файл формата .docx."), 400

    UPLOAD_DIR.mkdir(exist_ok=True)
    file.save(TEMPLATE_PATH)
    return jsonify(ok=True, message="✅ Шаблон акта успешно обновлён!")


@app.route("/admin/upload_offer", methods=["POST"])
def upload_offer_template():
    if "template" not in request.files:
        return jsonify(ok=False, error="Файл шаблона не передан."), 400

    file = request.files["template"]
    if not file or file.filename == "":
        return jsonify(ok=False, error="Файл не выбран."), 400

    if not file.filename.lower().endswith(".docx"):
        return jsonify(ok=False, error="Можно загрузить только файл формата .docx."), 400

    UPLOAD_DIR.mkdir(exist_ok=True)
    file.save(OFFER_TEMPLATE_PATH)
    return jsonify(ok=True, message="✅ Шаблон оферты успешно обновлён!")


@app.route("/generate", methods=["POST"])
def generate():
    payload = request.get_json(silent=True) or {}
    raw_acts = payload.get("acts", [])
    mode = payload.get("mode", "act_offer")  # act_offer | act_only | offer_only
    try:
        copies = int(payload.get("copies", 1))
    except (TypeError, ValueError):
        copies = 1
    copies = max(1, min(copies, 50))

    if not isinstance(raw_acts, list):
        return jsonify(ok=False, error="Некорректный формат данных."), 400

    acts = normalize_acts(raw_acts)
    if not acts:
        return jsonify(ok=False, error="Список актов пуст."), 400

    if mode in ("act_offer", "act_only") and not TEMPLATE_PATH.exists():
        return jsonify(ok=False, error="Шаблон акта uploads/template.docx не найден. Загрузите шаблон через Библиотеку шаблонов."), 400

    if mode in ("act_offer", "offer_only") and not OFFER_TEMPLATE_PATH.exists():
        return jsonify(ok=False, error="Шаблон оферты uploads/template_offer.docx не найден. Загрузите его через Библиотеку шаблонов."), 400

    try:
        zip_path = build_archive(acts, mode=mode, copies=copies)
    except Exception as exc:  # noqa: BLE001 - важно вернуть понятную ошибку пользователю
        return jsonify(ok=False, error=str(exc)), 500

    return jsonify(
        ok=True,
        archive=zip_path.name,
        download_url=url_for("download_zip", filename=zip_path.name),
        acts_count=len(acts),
        mode=mode,
    )


@app.route("/download/<path:filename>")
def download_zip(filename: str):
    # Защита от попыток выйти из папки output.
    if Path(filename).name != filename or not filename.endswith(".zip"):
        abort(404)

    zip_path = OUTPUT_DIR / filename
    if not zip_path.exists():
        abort(404)

    @after_this_request
    def remove_file_after_download(response):  # type: ignore[no-untyped-def]
        try:
            zip_path.unlink(missing_ok=True)
        except Exception:
            app.logger.exception("Не удалось удалить ZIP после скачивания: %s", zip_path)
        return response

    return send_file(
        zip_path,
        as_attachment=True,
        download_name=zip_path.name,
        mimetype="application/zip",
    )


@app.errorhandler(413)
def file_too_large(_error):  # type: ignore[no-untyped-def]
    return jsonify(ok=False, error="Файл слишком большой. Максимальный размер шаблона — 16MB."), 413


def normalize_acts(raw_acts: list[dict]) -> list[dict[str, str]]:
    """Оставляет только непустые строки и только поддерживаемые поля."""
    normalized: list[dict[str, str]] = []

    for raw in raw_acts:
        if not isinstance(raw, dict):
            continue

        row = {key: str(raw.get(key, "")).strip() for key in FIELD_KEYS}
        if any(row.values()):
            row["ORG_NAME"] = normalize_org_name(row.get("ORG_NAME", ""))
            if not row.get("DATE"):
                row["DATE"] = get_next_service_date(row.get("SIGN_DATETIME", ""))
            validity_dates = get_certificate_validity_dates(row.get("SIGN_DATETIME", ""))
            if validity_dates:
                row["VALID_FROM"], row["VALID_TO"] = validity_dates

            # Автосвязь дат оферты: OFFER_SENT_DATE и SIGN_DATETIMEOFF
            # ставятся за сутки до даты оказания услуг (DATE), если их
            # не ввели вручную.
            if row.get("DATE"):
                offer_sent, sign_off = get_offer_dates(row.get("DATE", ""))
                if offer_sent:
                    if not row.get("OFFER_SENT_DATE"):
                        row["OFFER_SENT_DATE"] = offer_sent
                    if not row.get("SIGN_DATETIMEOFF"):
                        row["SIGN_DATETIMEOFF"] = sign_off

            normalized.append(row)

    return normalized


def get_offer_dates(service_date: str) -> tuple[str, str]:
    """
    Даты оферты считаются от даты оказания услуг:
    OFFER_SENT_DATE и SIGN_DATETIMEOFF = DATE - 1 день.
    Возвращает (offer_sent, sign_off). Если DATE не распознана — ("", "").
    """
    parsed = try_parse_sign_datetime(service_date)
    if not parsed:
        return "", ""

    offer_day = parsed - timedelta(days=1)
    offer_sent = offer_day.strftime("%d.%m.%Y")
    # Подпись оферты — тот же день, время 00:00:00 (без времени) либо оставляем.
    sign_off = offer_day.strftime("%d.%m.%Y") + " 00:00:00"
    return offer_sent, sign_off


def normalize_org_name(org_name: str) -> str:
    """
    Исправляет частую ошибку ввода: 000/OOO вместо кириллического ООО.
    Например: 000 "Романс" -> ООО "Романс".
    """
    text = (org_name or "").strip()
    return re.sub(r"^[oOоО0]{3}(?=\s|[\"«]|$)", "ООО", text, count=1)


def get_next_service_date(sign_datetime: str) -> str:
    """Если дата услуг не заполнена, ставим следующий день после даты оферты/подписи."""
    parsed = try_parse_sign_datetime(sign_datetime)
    if not parsed:
        return ""
    return (parsed + timedelta(days=1)).strftime("%d.%m.%Y")


def get_certificate_validity_dates(sign_datetime: str) -> tuple[str, str] | None:
    """
    Даты сертификата считаются от даты подписи:
    VALID_FROM = дата подписи без времени,
    VALID_TO = та же дата + ровно 3 года.
    """
    parsed = try_parse_sign_datetime(sign_datetime)
    if not parsed:
        return None

    valid_from = parsed.strftime("%d.%m.%Y")
    valid_to = add_years(parsed, 3).strftime("%d.%m.%Y")
    return valid_from, valid_to


def add_years(value: datetime, years: int) -> datetime:
    try:
        return value.replace(year=value.year + years)
    except ValueError:
        # Например, 29.02.2024 + 3 года -> 28.02.2027.
        return value.replace(year=value.year + years, day=28)


def try_parse_sign_datetime(sign_datetime: str) -> datetime | None:
    value = (sign_datetime or "").strip()
    if not value:
        return None

    formats = [
        "%d.%m.%Y %H:%M:%S",
        "%d.%m.%Y %H:%M",
        "%d.%m.%Y",
        "%Y-%m-%dT%H:%M:%S",
        "%Y-%m-%dT%H:%M",
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%d %H:%M",
        "%Y-%m-%d",
    ]

    for date_format in formats:
        try:
            return datetime.strptime(value, date_format)
        except ValueError:
            continue

    return None


def build_archive(acts: list[dict[str, str]], mode: str = "act_offer", copies: int = 1) -> Path:
    OUTPUT_DIR.mkdir(exist_ok=True)

    run_dir = OUTPUT_DIR / f"run_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid4().hex[:8]}"
    run_dir.mkdir(parents=True, exist_ok=True)

    archive_name = f"acts_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
    archive_path = OUTPUT_DIR / archive_name

    generated_files: list[Path] = []
    used_base_names: set[str] = set()

    make_act = mode in ("act_offer", "act_only")
    make_offer = mode in ("act_offer", "offer_only")

    try:
        for act_number, act in enumerate(acts, start=1):
            replacements = prepare_replacements(act, act_number)

            for _copy in range(copies):
                if make_act:
                    base_name = make_act_filename_base(
                        sign_datetime=act.get("SIGN_DATETIME", ""),
                        act_number=act_number,
                        used_base_names=used_base_names,
                    )
                    docx_path = run_dir / f"{base_name}.docx"
                    pdf_path = run_dir / f"{base_name}.pdf"
                    create_docx_from_template(TEMPLATE_PATH, docx_path, replacements)
                    convert_docx_to_pdf(docx_path, pdf_path)
                    generated_files.extend([docx_path, pdf_path])

                if make_offer:
                    base_name = make_offer_filename_base(
                        sign_datetime=act.get("SIGN_DATETIMEOFF", "") or act.get("SIGN_DATETIME", ""),
                        act_number=act_number,
                        used_base_names=used_base_names,
                    )
                    docx_path = run_dir / f"{base_name}.docx"
                    pdf_path = run_dir / f"{base_name}.pdf"
                    create_docx_from_template(OFFER_TEMPLATE_PATH, docx_path, replacements)
                    convert_docx_to_pdf(docx_path, pdf_path)
                    generated_files.extend([docx_path, pdf_path])

        with zipfile.ZipFile(archive_path, "w", compression=zipfile.ZIP_DEFLATED) as zip_file:
            for file_path in generated_files:
                zip_file.write(file_path, arcname=file_path.name)

        return archive_path

    finally:
        # DOCX/PDF удаляются сразу после упаковки в ZIP или после ошибки.
        shutil.rmtree(run_dir, ignore_errors=True)


def make_act_filename_base(sign_datetime: str, act_number: int, used_base_names: set[str]) -> str:
    """
    Формирует имя файла для акта:
    Акт оказанных услуг - {SIGN_DATETIME_SHORT}-{ACT_DATE}.{RND2}
    """
    parsed_datetime = parse_sign_datetime_for_filename(sign_datetime, act_number)
    sign_datetime_short = parsed_datetime.strftime("%Y%m%d%H%M%S")
    act_date = parsed_datetime.strftime("%Y-%m-%d")

    for _attempt in range(100):
        rnd2 = f"{secrets.randbelow(100):02d}"
        base_name = f"Акт оказанных услуг - {sign_datetime_short}-{act_date}.{rnd2}"
        if base_name not in used_base_names:
            used_base_names.add(base_name)
            return base_name

    raise RuntimeError(
        f"Не удалось подобрать уникальное имя файла для строки №{act_number}. "
        "Измените дату/время подписи или повторите генерацию."
    )


def make_offer_filename_base(sign_datetime: str, act_number: int, used_base_names: set[str]) -> str:
    """
    Формирует имя файла по новому правилу:
    Оферта - {SIGN_DATETIME_SHORT}-{OFFER_DATE}.3{RND2}
    """
    parsed_datetime = parse_sign_datetime_for_filename(sign_datetime, act_number)
    sign_datetime_short = parsed_datetime.strftime("%Y%m%d%H%M%S")
    offer_date = parsed_datetime.strftime("%Y-%m-%d")

    for _attempt in range(100):
        rnd2 = f"{secrets.randbelow(100):02d}"
        base_name = f"Оферта - {sign_datetime_short}-{offer_date}.3{rnd2}"
        if base_name not in used_base_names:
            used_base_names.add(base_name)
            return base_name

    raise RuntimeError(
        f"Не удалось подобрать уникальное имя файла для строки №{act_number}. "
        "Измените дату/время подписи или повторите генерацию."
    )


def parse_sign_datetime_for_filename(sign_datetime: str, act_number: int) -> datetime:
    value = (sign_datetime or "").strip()
    if not value:
        raise RuntimeError(
            f"В строке №{act_number} заполните поле «Дата подписи». "
            "Оно используется в имени файла."
        )

    parsed = try_parse_sign_datetime(value)
    if parsed:
        return parsed

    raise RuntimeError(
        f"В строке №{act_number} поле «Дата подписи» должно быть в формате "
        "ДД.ММ.ГГГГ ЧЧ:ММ:СС. Например: 21.06.2024 14:37:52."
    )


def prepare_replacements(act: dict[str, str], act_number: int) -> dict[str, str]:
    replacements = {key: act.get(key, "") for key in FIELD_KEYS}
    replacements["ACT_NUMBER"] = str(act_number)
    replacements["CERT_ORG"] = make_certificate_number()
    replacements["CERT_EXEC"] = make_certificate_number()
    replacements["CERT_ORGOFF"] = make_certificate_number()
    replacements["CERT_ORGOFF_ORG"] = make_certificate_number()
    return replacements


def make_certificate_number() -> str:
    """Случайный номер сертификата в формате UUID: 8-4-4-4-12."""
    return str(uuid4())


def create_docx_from_template(template_path: Path, output_path: Path, replacements: dict[str, str]) -> None:
    document = Document(template_path)
    replace_tags_in_document(document, replacements)
    # Убираем колонтитулы в самом конце, чтобы они не сдвигали содержимое
    # вниз/вверх. Важно делать ПОСЛЕ подстановки меток: обращение к
    # section.header/section.footer в iter_text_xml_roots создаёт колонтитулы
    # заново.
    remove_headers_footers(document)
    document.save(output_path)


def remove_headers_footers(document) -> None:  # type: ignore[no-untyped-def]
    """Удаляет ссылки на верхние и нижние колонтитулы во всех секциях."""
    from docx.oxml.ns import qn

    for section in document.sections:
        sect_pr = section._sectPr
        if sect_pr is None:
            continue
        for tag in ("w:headerReference", "w:footerReference"):
            for ref in sect_pr.findall(qn(tag)):
                sect_pr.remove(ref)
        # Убираем признак отдельного колонтитула первой страницы.
        title_pg = sect_pr.find(qn("w:titlePg"))
        if title_pg is not None:
            sect_pr.remove(title_pg)


def replace_tags_in_document(document, replacements: dict[str, str]) -> None:  # type: ignore[no-untyped-def]
    tags = {f"{{{{{key}}}}}": value for key, value in replacements.items()}

    # Поддержка старых шаблонов, где сертификаты прописаны не метками, а
    # фиксированным текстом. Меняем их на новые случайные номера.
    if replacements.get("CERT_ORG"):
        tags[DEFAULT_ORG_CERT] = replacements["CERT_ORG"]
    if replacements.get("CERT_EXEC"):
        tags[DEFAULT_EXEC_CERT] = replacements["CERT_EXEC"]

    # Обычные параграфы и ячейки таблиц.
    for paragraph in iter_all_paragraphs(document):
        replace_tags_in_paragraph(paragraph, tags)

    # Дополнительно проходим по XML-узлам текста. Это помогает заменить метки
    # внутри текстовых блоков/фигур Word, если шаблон использует такие элементы
    # для визуального оформления электронной подписи.
    for root in iter_text_xml_roots(document):
        replace_tags_in_xml_text_nodes(root, tags)


def iter_text_xml_roots(document):  # type: ignore[no-untyped-def]
    """XML-корни, где могут находиться w:t, включая текстовые блоки Word."""
    yielded_ids: set[int] = set()

    roots = [document._element]
    for section in document.sections:
        roots.extend(
            [
                section.header._element,
                section.footer._element,
                section.first_page_header._element,
                section.first_page_footer._element,
                section.even_page_header._element,
                section.even_page_footer._element,
            ]
        )

    for root in roots:
        root_id = id(root)
        if root_id not in yielded_ids:
            yielded_ids.add(root_id)
            yield root


def replace_tags_in_xml_text_nodes(root, tags: dict[str, str]) -> None:  # type: ignore[no-untyped-def]
    """
    Заменяет метки по всем XML-узлам w:t.

    Нужна для текстовых блоков/фигур Word: python-docx не показывает их как
    обычные paragraph/table, но сами текстовые узлы в XML доступны.
    """
    text_nodes = list(root.iter(qn("w:t")))
    if not text_nodes:
        return

    for tag, value in tags.items():
        while True:
            full_text = "".join(node.text or "" for node in text_nodes)
            start = full_text.find(tag)
            if start == -1:
                break

            end = start + len(tag)
            char_map: list[tuple[int, int]] = []
            for node_index, node in enumerate(text_nodes):
                for char_offset, _char in enumerate(node.text or ""):
                    char_map.append((node_index, char_offset))

            if end > len(char_map):
                break

            start_node_index, start_offset = char_map[start]
            end_node_index, end_offset = char_map[end - 1]

            first_node = text_nodes[start_node_index]
            last_node = text_nodes[end_node_index]
            before = (first_node.text or "")[:start_offset]
            after = (last_node.text or "")[end_offset + 1 :]
            first_node.text = before + str(value) + after

            for node_index in range(start_node_index + 1, end_node_index + 1):
                text_nodes[node_index].text = ""


def iter_all_paragraphs(document):  # type: ignore[no-untyped-def]
    """Параграфы документа, таблиц, а также колонтитулов."""
    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        yield from iter_table_paragraphs(table)

    for section in document.sections:
        header_footer_parts = [
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ]
        for part in header_footer_parts:
            for paragraph in part.paragraphs:
                yield paragraph
            for table in part.tables:
                yield from iter_table_paragraphs(table)


def iter_table_paragraphs(table):  # type: ignore[no-untyped-def]
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                yield from iter_table_paragraphs(nested_table)


def replace_tags_in_paragraph(paragraph, tags: dict[str, str]) -> None:  # type: ignore[no-untyped-def]
    """
    Заменяет метки даже если Word разбил {{TAG}} на несколько runs.
    Форматирование берётся из первого run, в котором начинается метка.
    """
    if not paragraph.runs:
        return

    for tag, value in tags.items():
        while True:
            full_text = "".join(run.text for run in paragraph.runs)
            start = full_text.find(tag)
            if start == -1:
                break

            end = start + len(tag)
            char_map: list[tuple[int, int]] = []
            for run_index, run in enumerate(paragraph.runs):
                for char_offset, _char in enumerate(run.text):
                    char_map.append((run_index, char_offset))

            if end > len(char_map):
                break

            start_run_index, start_offset = char_map[start]
            end_run_index, end_offset = char_map[end - 1]

            runs = paragraph.runs
            first_run = runs[start_run_index]
            last_run = runs[end_run_index]

            before = first_run.text[:start_offset]
            after = last_run.text[end_offset + 1 :]
            first_run.text = before + str(value) + after

            for run_index in range(start_run_index + 1, end_run_index + 1):
                runs[run_index].text = ""


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    """
    Конвертация DOCX -> PDF через docx2pdf.

    На Windows docx2pdf управляет Microsoft Word через COM. Flask обрабатывает
    запросы в отдельных потоках, поэтому COM нужно явно инициализировать через
    pythoncom.CoInitialize(), иначе появляется ошибка:
    "Не был произведен вызов CoInitialize".
    """
    pythoncom_module = None
    com_initialized = False

    try:
        if os.name == "nt":
            import pythoncom

            pythoncom_module = pythoncom
            pythoncom_module.CoInitialize()
            com_initialized = True

        from docx2pdf import convert

        # Word/COM плохо переносит параллельные обращения, поэтому конвертируем
        # PDF последовательно даже если на сервер пришло несколько запросов.
        with PDF_CONVERSION_LOCK:
            convert(str(docx_path), str(pdf_path))

    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(
            "PDF не генерируется. Убедитесь, что на сервере установлен Microsoft Word, "
            "Word запускается без всплывающих окон/активации, и установлена библиотека "
            "docx2pdf. Подробности: " + str(exc)
        ) from exc

    finally:
        if com_initialized and pythoncom_module is not None:
            try:
                pythoncom_module.CoUninitialize()
            except Exception:
                pass

    if not pdf_path.exists():
        raise RuntimeError("PDF не был создан. Проверьте установку Microsoft Word и docx2pdf.")


def safe_filename_part(value: str | None, default: str) -> str:
    text = (value or default).strip() or default
    text = re.sub(r"\s+", "_", text)
    text = re.sub(r"[<>:\"/\\|?*\x00-\x1f]+", "_", text)
    text = text.strip("._ ")
    return (text or default)[:120]


def get_offer_date_part(sign_datetime: str) -> str:
    raw_date = (sign_datetime or "")[:10].strip()
    if not raw_date:
        return "Без_даты"

    safe_date = raw_date.replace(".", "-").replace("/", "-").replace("\\", "-")
    return safe_filename_part(safe_date, default="Без_даты")


if __name__ == "__main__":
    port = int(os.environ.get("PORT", "5000"))
    app.run(host="0.0.0.0", port=port, debug=True)
